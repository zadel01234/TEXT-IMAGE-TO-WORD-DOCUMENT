import os
import threading
import tkinter as tk
from tkinter import filedialog, ttk, messagebox
from PIL import Image, ImageTk
import pytesseract
from docx import Document

class ImgTextToWordGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Image text to Word Converter")
        self.root.geometry("800x600")
        self.root.minsize(700, 500)
        
        # Set up variables
        self.input_image_path = tk.StringVar()
        self.output_doc_path = tk.StringVar()
        self.language = tk.StringVar(value="eng")
        self.status_message = tk.StringVar(value="Ready")
        self.preview_image = None
        
        # Available OCR languages - common ones
        self.languages = {
            "English": "eng",
            "Spanish": "spa",
            "French": "fra",
            "German": "deu",
            "Chinese (Simplified)": "chi_sim",
            "Chinese (Traditional)": "chi_tra",
            "Japanese": "jpn",
            "Korean": "kor",
            "Russian": "rus",
            "Arabic": "ara"
        }
        
        # Create the main frame
        self.main_frame = ttk.Frame(root, padding="10")
        self.main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Create the UI elements
        self.create_input_section()
        self.create_preview_section()
        self.create_output_section()
        self.create_process_section()
        self.create_status_bar()

    def create_input_section(self):
        input_frame = ttk.LabelFrame(self.main_frame, text="Image Input", padding="10")
        input_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(input_frame, text="Image File:").grid(row=0, column=0, sticky=tk.W, pady=5)
        ttk.Entry(input_frame, textvariable=self.input_image_path, width=50).grid(row=0, column=1, sticky=tk.EW, pady=5, padx=5)
        ttk.Button(input_frame, text="Browse...", command=self.browse_input_image).grid(row=0, column=2, pady=5)
        
        ttk.Label(input_frame, text="Language:").grid(row=1, column=0, sticky=tk.W, pady=5)
        language_combo = ttk.Combobox(input_frame, textvariable=self.language, state="readonly")
        language_combo['values'] = list(self.languages.keys())
        language_combo.current(0)  # Default to English
        language_combo.grid(row=1, column=1, sticky=tk.W, pady=5, padx=5)
        language_combo.bind('<<ComboboxSelected>>', self.on_language_selected)

    def create_preview_section(self):
        preview_frame = ttk.LabelFrame(self.main_frame, text="Image Preview", padding="10")
        preview_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        self.canvas = tk.Canvas(preview_frame, bg="white", highlightthickness=1, highlightbackground="gray")
        self.canvas.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

    def create_output_section(self):
        output_frame = ttk.LabelFrame(self.main_frame, text="Word Output", padding="10")
        output_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(output_frame, text="Output File:").grid(row=0, column=0, sticky=tk.W, pady=5)
        ttk.Entry(output_frame, textvariable=self.output_doc_path, width=50).grid(row=0, column=1, sticky=tk.EW, pady=5, padx=5)
        ttk.Button(output_frame, text="Browse...", command=self.browse_output_file).grid(row=0, column=2, pady=5)

    def create_process_section(self):
        process_frame = ttk.Frame(self.main_frame)
        process_frame.pack(fill=tk.X, padx=5, pady=10)
        
        self.progress_bar = ttk.Progressbar(process_frame, orient=tk.HORIZONTAL, length=100, mode='indeterminate')
        self.progress_bar.pack(fill=tk.X, pady=5)
        
        button_frame = ttk.Frame(process_frame)
        button_frame.pack(fill=tk.X, pady=5)
        
        ttk.Button(button_frame, text="Convert to Word", command=self.process_image).pack(side=tk.RIGHT, padx=5)

    def create_status_bar(self):
        status_bar = ttk.Frame(self.root, relief=tk.SUNKEN, padding=(2, 2))
        status_bar.pack(side=tk.BOTTOM, fill=tk.X)
        
        ttk.Label(status_bar, textvariable=self.status_message).pack(side=tk.LEFT)

    def browse_input_image(self):
        filetypes = (
            ('Image files', '*.png *.jpg *.jpeg *.bmp *.tiff *.tif'),
            ('All files', '*.*')
        )
        filename = filedialog.askopenfilename(
            title='Open an image file',
            initialdir='/',
            filetypes=filetypes
        )
        
        if filename:
            self.input_image_path.set(filename)
            self.load_preview_image(filename)
            
            # Auto-generate output filename
            base_name = os.path.splitext(os.path.basename(filename))[0]
            self.output_doc_path.set(os.path.join(os.path.dirname(filename), f"{base_name}.docx"))

    def load_preview_image(self, image_path):
        try:
            # Open the image
            img = Image.open(image_path)
            
            # Resize to fit canvas while maintaining aspect ratio
            canvas_width = self.canvas.winfo_width() or 700
            canvas_height = self.canvas.winfo_height() or 300
            
            img.thumbnail((canvas_width, canvas_height))
            
            # Convert to PhotoImage and keep a reference
            self.preview_image = ImageTk.PhotoImage(img)
            
            # Clear canvas and display image
            self.canvas.delete("all")
            self.canvas.create_image(canvas_width//2, canvas_height//2, image=self.preview_image, anchor=tk.CENTER)
            
            self.status_message.set(f"Loaded image: {os.path.basename(image_path)}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load image: {str(e)}")
            self.status_message.set("Error loading image")

    def browse_output_file(self):
        filetypes = (
            ('Word documents', '*.docx'),
            ('All files', '*.*')
        )
        filename = filedialog.asksaveasfilename(
            title='Save as Word document',
            initialdir='/',
            defaultextension=".docx",
            filetypes=filetypes
        )
        
        if filename:
            self.output_doc_path.set(filename)

    def on_language_selected(self, event=None):
        language_name = event.widget.get()
        self.language.set(self.languages[language_name])

    def process_image(self):
        # Validate input
        if not self.input_image_path.get():
            messagebox.showerror("Error", "Please select an input image.")
            return
        
        if not self.output_doc_path.get():
            messagebox.showerror("Error", "Please specify an output document path.")
            return
        
        # Start processing in a separate thread to keep UI responsive
        self.progress_bar.start()
        self.status_message.set("Processing...")
        threading.Thread(target=self.ocr_to_word_thread, daemon=True).start()

    def ocr_to_word_thread(self):
        try:
            image_path = self.input_image_path.get()
            output_file = self.output_doc_path.get()
            lang = self.language.get()
            
            # Open the image
            image = Image.open(image_path)
            
            # Extract text using pytesseract OCR
            text = pytesseract.image_to_string(image, lang=lang)
            
            # Create a new Word document
            doc = Document()
            
            # Add a title
            doc.add_heading('Extracted Text From Image', 0)
            
            # Process the text content
            paragraphs = text.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
            
            # Save the document
            doc.save(output_file)
            
            # Update UI on the main thread
            self.root.after(0, self.process_complete, True, f"Document successfully created: {os.path.basename(output_file)}")
            
        except Exception as e:
            self.root.after(0, self.process_complete, False, str(e))

    def process_complete(self, success, message):
        self.progress_bar.stop()
        
        if success:
            self.status_message.set(message)
            messagebox.showinfo("Success", message)
            
            # Ask if user wants to open the document
            if messagebox.askyesno("Open Document", "Would you like to open the created document?"):
                self.open_document(self.output_doc_path.get())
        else:
            self.status_message.set(f"Error: {message}")
            messagebox.showerror("Error", f"Failed to convert image: {message}")

    def open_document(self, doc_path):
        try:
            import platform
            import subprocess
            
            if platform.system() == 'Windows':
                os.startfile(doc_path)
            elif platform.system() == 'Darwin':  # macOS
                subprocess.call(('open', doc_path))
            else:  # Linux
                subprocess.call(('xdg-open', doc_path))
        except Exception as e:
            messagebox.showerror("Error", f"Could not open document: {str(e)}")

def main():
    root = tk.Tk()
    app = ImgTextToWordGUI(root)
    root.mainloop()

if __name__ == "__main__":
    main()
