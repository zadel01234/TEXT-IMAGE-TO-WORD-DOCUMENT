import os
import threading
import tkinter as tk
from tkinter import filedialog, ttk, messagebox, scrolledtext
from PIL import Image, ImageTk, ImageEnhance, ImageOps
import pytesseract
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import math


class OCRtoWordGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Advanced OCR to Word Converter")
        self.root.geometry("1000x700")
        self.root.minsize(900, 600)

        # Set up variables
        self.input_image_path = tk.StringVar()
        self.output_doc_path = tk.StringVar()
        self.language = tk.StringVar(value="eng")
        self.status_message = tk.StringVar(value="Ready")
        self.preview_image = None
        self.original_image = None
        self.preview_scale = 1.0
        self.rotation_angle = 0

        # Batch processing variables
        self.batch_files = []
        self.batch_mode = tk.BooleanVar(value=False)

        # Document formatting variables
        self.font_family = tk.StringVar(value="Calibri")
        self.font_size = tk.IntVar(value=11)
        self.alignment = tk.StringVar(value="Left")
        self.include_title = tk.BooleanVar(value=True)
        self.title_text = tk.StringVar(value="OCR Extracted Text")

        # Image processing variables
        self.brightness = tk.DoubleVar(value=1.0)
        self.contrast = tk.DoubleVar(value=1.0)
        self.sharpen = tk.DoubleVar(value=1.0)
        self.binarize = tk.BooleanVar(value=False)
        self.threshold = tk.IntVar(value=127)

        # Available OCR languages
        self.languages = {
            "English": "eng",
            "Spanish": "spa",
            "French": "fra",
            "German": "deu",
            "Chinese (Simplified)": "chi_sim",
            "Chinese (Traditional)": "chi_tra",
            "Japanese": "jpn",
            "Korean": "kor",
            "Russian": "rus",
            "Arabic": "ara",
            "Hindi": "hin",
            "Italian": "ita",
            "Portuguese": "por",
            "Dutch": "nld",
            "Turkish": "tur",
            "Hebrew": "heb",
            "Polish": "pol",
            "Czech": "ces",
            "Greek": "ell",
            "Thai": "tha"
        }

        # Font families
        self.font_families = [
            "Calibri", "Arial", "Times New Roman", "Courier New",
            "Verdana", "Tahoma", "Georgia", "Garamond", "Comic Sans MS"
        ]

        # Create the notebook for tabbed interface
        self.notebook = ttk.Notebook(root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Create main tabs
        self.main_tab = ttk.Frame(self.notebook)
        self.batch_tab = ttk.Frame(self.notebook)
        self.format_tab = ttk.Frame(self.notebook)
        self.process_tab = ttk.Frame(self.notebook)

        self.notebook.add(self.main_tab, text="Main")
        self.notebook.add(self.batch_tab, text="Batch Processing")
        self.notebook.add(self.format_tab, text="Document Format")
        self.notebook.add(self.process_tab, text="Image Processing")

        # Create the UI elements
        self.create_main_tab()
        self.create_batch_tab()
        self.create_format_tab()
        self.create_process_tab()
        self.create_status_bar()

        # Create the extracted text preview window (initially hidden)
        self.text_preview_window = None

    def create_main_tab(self):
        main_frame = ttk.Frame(self.main_tab, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # Input section
        input_frame = ttk.LabelFrame(main_frame, text="Image Input", padding="10")
        input_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Label(input_frame, text="Image File:").grid(row=0, column=0, sticky=tk.W, pady=5)
        ttk.Entry(input_frame, textvariable=self.input_image_path, width=50).grid(row=0, column=1, sticky=tk.EW, pady=5,
                                                                                  padx=5)
        ttk.Button(input_frame, text="Browse...", command=self.browse_input_image).grid(row=0, column=2, pady=5)

        ttk.Label(input_frame, text="Language:").grid(row=1, column=0, sticky=tk.W, pady=5)
        language_combo = ttk.Combobox(input_frame, textvariable=self.language, state="readonly")
        language_combo['values'] = list(self.languages.keys())
        language_combo.current(0)
        language_combo.grid(row=1, column=1, sticky=tk.W, pady=5, padx=5)
        language_combo.bind('<<ComboboxSelected>>', self.on_language_selected)

        # Preview section with zoom and rotate controls
        preview_frame = ttk.LabelFrame(main_frame, text="Image Preview", padding="10")
        preview_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Canvas for image preview
        self.canvas_frame = ttk.Frame(preview_frame)
        self.canvas_frame.pack(fill=tk.BOTH, expand=True, side=tk.LEFT)

        self.canvas = tk.Canvas(self.canvas_frame, bg="white", highlightthickness=1, highlightbackground="gray")
        self.canvas.pack(fill=tk.BOTH, expand=True)

        # Scrollbars for canvas
        h_scrollbar = ttk.Scrollbar(self.canvas_frame, orient=tk.HORIZONTAL, command=self.canvas.xview)
        v_scrollbar = ttk.Scrollbar(self.canvas_frame, orient=tk.VERTICAL, command=self.canvas.yview)
        h_scrollbar.pack(fill=tk.X, side=tk.BOTTOM)
        v_scrollbar.pack(fill=tk.Y, side=tk.RIGHT)

        self.canvas.configure(xscrollcommand=h_scrollbar.set, yscrollcommand=v_scrollbar.set)
        self.canvas.bind("<ButtonPress-1>", self.scroll_start)
        self.canvas.bind("<B1-Motion>", self.scroll_move)

        # Controls for the image
        controls_frame = ttk.Frame(preview_frame)
        controls_frame.pack(fill=tk.Y, side=tk.RIGHT, padx=5)

        ttk.Label(controls_frame, text="Zoom:").pack(anchor=tk.W, pady=(0, 5))
        zoom_frame = ttk.Frame(controls_frame)
        zoom_frame.pack(fill=tk.X, pady=(0, 10))
        ttk.Button(zoom_frame, text="-", width=3, command=self.zoom_out).pack(side=tk.LEFT)
        ttk.Button(zoom_frame, text="+", width=3, command=self.zoom_in).pack(side=tk.RIGHT)
        ttk.Button(zoom_frame, text="Fit", width=5, command=self.zoom_fit).pack(side=tk.LEFT, padx=5)
        ttk.Button(zoom_frame, text="100%", width=5, command=self.zoom_reset).pack(side=tk.RIGHT, padx=5)

        ttk.Label(controls_frame, text="Rotate:").pack(anchor=tk.W, pady=(0, 5))
        rotate_frame = ttk.Frame(controls_frame)
        rotate_frame.pack(fill=tk.X, pady=(0, 10))
        ttk.Button(rotate_frame, text="↶", width=3, command=self.rotate_ccw).pack(side=tk.LEFT)
        ttk.Button(rotate_frame, text="↷", width=3, command=self.rotate_cw).pack(side=tk.RIGHT)
        ttk.Button(rotate_frame, text="Reset", command=self.rotate_reset).pack(side=tk.LEFT, padx=5, fill=tk.X,
                                                                               expand=True)

        # Output section
        output_frame = ttk.LabelFrame(main_frame, text="Word Output", padding="10")
        output_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Label(output_frame, text="Output File:").grid(row=0, column=0, sticky=tk.W, pady=5)
        ttk.Entry(output_frame, textvariable=self.output_doc_path, width=50).grid(row=0, column=1, sticky=tk.EW, pady=5,
                                                                                  padx=5)
        ttk.Button(output_frame, text="Browse...", command=self.browse_output_file).grid(row=0, column=2, pady=5)

        # Process buttons
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X, pady=10)

        self.progress_bar = ttk.Progressbar(button_frame, orient=tk.HORIZONTAL, length=100, mode='indeterminate')
        self.progress_bar.pack(fill=tk.X, pady=5)

        actions_frame = ttk.Frame(button_frame)
        actions_frame.pack(fill=tk.X)

        ttk.Button(actions_frame, text="Preview Extracted Text", command=self.preview_text).pack(side=tk.LEFT, padx=5)
        ttk.Button(actions_frame, text="Apply Image Processing", command=self.apply_image_processing).pack(side=tk.LEFT,
                                                                                                           padx=5)
        ttk.Button(actions_frame, text="Convert to Word", command=self.process_image).pack(side=tk.RIGHT, padx=5)

    def create_batch_tab(self):
        batch_frame = ttk.Frame(self.batch_tab, padding="10")
        batch_frame.pack(fill=tk.BOTH, expand=True)

        # Batch mode checkbox
        ttk.Checkbutton(batch_frame, text="Enable Batch Processing", variable=self.batch_mode).pack(anchor=tk.W, pady=5)

        # Files selection section
        files_frame = ttk.LabelFrame(batch_frame, text="Batch Files", padding="10")
        files_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        buttons_frame = ttk.Frame(files_frame)
        buttons_frame.pack(fill=tk.X, pady=5)
        ttk.Button(buttons_frame, text="Add Files", command=self.add_batch_files).pack(side=tk.LEFT, padx=5)
        ttk.Button(buttons_frame, text="Remove Selected", command=self.remove_selected_files).pack(side=tk.LEFT, padx=5)
        ttk.Button(buttons_frame, text="Clear All", command=self.clear_batch_files).pack(side=tk.LEFT, padx=5)

        # Files listbox with scrollbar
        list_frame = ttk.Frame(files_frame)
        list_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        self.files_listbox = tk.Listbox(list_frame, selectmode=tk.EXTENDED)
        self.files_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=self.files_listbox.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.files_listbox.configure(yscrollcommand=scrollbar.set)

        # Output directory selection
        output_frame = ttk.LabelFrame(batch_frame, text="Batch Output Directory", padding="10")
        output_frame.pack(fill=tk.X, padx=5, pady=5)

        self.batch_output_dir = tk.StringVar()
        ttk.Entry(output_frame, textvariable=self.batch_output_dir, width=50).pack(side=tk.LEFT, fill=tk.X, expand=True,
                                                                                   padx=5)
        ttk.Button(output_frame, text="Browse...", command=self.browse_output_dir).pack(side=tk.RIGHT, padx=5)

        # Process batch button
        ttk.Button(batch_frame, text="Process Batch", command=self.process_batch).pack(side=tk.RIGHT, pady=10, padx=5)

    def create_format_tab(self):
        format_frame = ttk.Frame(self.format_tab, padding="10")
        format_frame.pack(fill=tk.BOTH, expand=True)

        # Text formatting options
        text_format_frame = ttk.LabelFrame(format_frame, text="Text Formatting", padding="10")
        text_format_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Label(text_format_frame, text="Font Family:").grid(row=0, column=0, sticky=tk.W, pady=5)
        font_combo = ttk.Combobox(text_format_frame, textvariable=self.font_family, state="readonly")
        font_combo['values'] = self.font_families
        font_combo.grid(row=0, column=1, sticky=tk.W, pady=5, padx=5)

        ttk.Label(text_format_frame, text="Font Size:").grid(row=1, column=0, sticky=tk.W, pady=5)
        size_frame = ttk.Frame(text_format_frame)
        size_frame.grid(row=1, column=1, sticky=tk.W, pady=5, padx=5)
        sizes = [8, 9, 10, 11, 12, 14, 16, 18, 20, 24, 28, 32, 36]
        size_combo = ttk.Combobox(size_frame, textvariable=self.font_size, state="readonly", width=5)
        size_combo['values'] = sizes
        size_combo.pack(side=tk.LEFT)

        ttk.Label(text_format_frame, text="Alignment:").grid(row=2, column=0, sticky=tk.W, pady=5)
        alignment_frame = ttk.Frame(text_format_frame)
        alignment_frame.grid(row=2, column=1, sticky=tk.W, pady=5, padx=5)
        ttk.Radiobutton(alignment_frame, text="Left", variable=self.alignment, value="Left").pack(side=tk.LEFT, padx=5)
        ttk.Radiobutton(alignment_frame, text="Center", variable=self.alignment, value="Center").pack(side=tk.LEFT,
                                                                                                      padx=5)
        ttk.Radiobutton(alignment_frame, text="Right", variable=self.alignment, value="Right").pack(side=tk.LEFT,
                                                                                                    padx=5)
        ttk.Radiobutton(alignment_frame, text="Justify", variable=self.alignment, value="Justify").pack(side=tk.LEFT,
                                                                                                        padx=5)

        # Title options
        title_frame = ttk.LabelFrame(format_frame, text="Document Title", padding="10")
        title_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Checkbutton(title_frame, text="Include title in document", variable=self.include_title).grid(row=0,
                                                                                                         column=0,
                                                                                                         sticky=tk.W,
                                                                                                         pady=5)

        ttk.Label(title_frame, text="Title text:").grid(row=1, column=0, sticky=tk.W, pady=5)
        ttk.Entry(title_frame, textvariable=self.title_text, width=50).grid(row=1, column=1, sticky=tk.EW, pady=5,
                                                                            padx=5)

        # Document preview section
        preview_frame = ttk.LabelFrame(format_frame, text="Document Preview", padding="10")
        preview_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # This would show a mockup of the document format
        ttk.Label(preview_frame, text="[Document preview not available in this version]").pack(pady=20)

    def create_process_tab(self):
        process_frame = ttk.Frame(self.process_tab, padding="10")
        process_frame.pack(fill=tk.BOTH, expand=True)

        # Image processing options
        adjustment_frame = ttk.LabelFrame(process_frame, text="Image Adjustments", padding="10")
        adjustment_frame.pack(fill=tk.X, padx=5, pady=5)

        # Brightness control
        ttk.Label(adjustment_frame, text="Brightness:").grid(row=0, column=0, sticky=tk.W, pady=5)
        brightness_scale = ttk.Scale(adjustment_frame, from_=0.5, to=2.0, orient=tk.HORIZONTAL,
                                     variable=self.brightness, length=200)
        brightness_scale.grid(row=0, column=1, sticky=tk.EW, pady=5, padx=5)
        ttk.Label(adjustment_frame, textvariable=tk.StringVar(value=lambda: f"{self.brightness.get():.1f}")).grid(row=0,
                                                                                                                  column=2,
                                                                                                                  padx=5)

        # Contrast control
        ttk.Label(adjustment_frame, text="Contrast:").grid(row=1, column=0, sticky=tk.W, pady=5)
        contrast_scale = ttk.Scale(adjustment_frame, from_=0.5, to=2.0, orient=tk.HORIZONTAL,
                                   variable=self.contrast, length=200)
        contrast_scale.grid(row=1, column=1, sticky=tk.EW, pady=5, padx=5)
        ttk.Label(adjustment_frame, textvariable=tk.StringVar(value=lambda: f"{self.contrast.get():.1f}")).grid(row=1,
                                                                                                                column=2,
                                                                                                                padx=5)

        # Sharpness control
        ttk.Label(adjustment_frame, text="Sharpness:").grid(row=2, column=0, sticky=tk.W, pady=5)
        sharpen_scale = ttk.Scale(adjustment_frame, from_=0.0, to=2.0, orient=tk.HORIZONTAL,
                                  variable=self.sharpen, length=200)
        sharpen_scale.grid(row=2, column=1, sticky=tk.EW, pady=5, padx=5)
        ttk.Label(adjustment_frame, textvariable=tk.StringVar(value=lambda: f"{self.sharpen.get():.1f}")).grid(row=2,
                                                                                                               column=2,
                                                                                                               padx=5)

        # Binarization controls
        binary_frame = ttk.LabelFrame(process_frame, text="Binarization (Black & White)", padding="10")
        binary_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Checkbutton(binary_frame, text="Convert to black and white", variable=self.binarize).grid(row=0, column=0,
                                                                                                      sticky=tk.W,
                                                                                                      pady=5)

        ttk.Label(binary_frame, text="Threshold:").grid(row=1, column=0, sticky=tk.W, pady=5)
        threshold_scale = ttk.Scale(binary_frame, from_=0, to=255, orient=tk.HORIZONTAL,
                                    variable=self.threshold, length=200)
        threshold_scale.grid(row=1, column=1, sticky=tk.EW, pady=5, padx=5)
        ttk.Label(binary_frame, textvariable=self.threshold).grid(row=1, column=2, padx=5)

        # Apply buttons
        button_frame = ttk.Frame(process_frame)
        button_frame.pack(fill=tk.X, pady=10)

        ttk.Button(button_frame, text="Reset to Default", command=self.reset_image_processing).pack(side=tk.LEFT,
                                                                                                    padx=5)
        ttk.Button(button_frame, text="Apply to Current Image", command=self.apply_image_processing).pack(side=tk.RIGHT,
                                                                                                          padx=5)

    def create_status_bar(self):
        status_bar = ttk.Frame(self.root, relief=tk.SUNKEN, padding=(2, 2))
        status_bar.pack(side=tk.BOTTOM, fill=tk.X)

        ttk.Label(status_bar, textvariable=self.status_message).pack(side=tk.LEFT)

        # Add zoom level indicator
        self.zoom_info = tk.StringVar(value="Zoom: 100%")
        ttk.Label(status_bar, textvariable=self.zoom_info).pack(side=tk.RIGHT, padx=10)

    # Image preview interaction methods
    def scroll_start(self, event):
        self.canvas.scan_mark(event.x, event.y)

    def scroll_move(self, event):
        self.canvas.scan_dragto(event.x, event.y, gain=1)

    def zoom_in(self):
        self.preview_scale *= 1.25
        self.update_preview()

    def zoom_out(self):
        self.preview_scale *= 0.8
        self.update_preview()

    def zoom_reset(self):
        self.preview_scale = 1.0
        self.update_preview()

    def zoom_fit(self):
        if self.original_image:
            canvas_width = self.canvas.winfo_width()
            canvas_height = self.canvas.winfo_height()
            if canvas_width <= 1:  # Not yet fully initialized
                canvas_width = 600
                canvas_height = 400

            img_width, img_height = self.original_image.size

            # Calculate scale to fit
            scale_w = canvas_width / img_width
            scale_h = canvas_height / img_height
            self.preview_scale = min(scale_w, scale_h) * 0.9  # 90% of fit

            self.update_preview()

    def rotate_cw(self):
        self.rotation_angle = (self.rotation_angle + 90) % 360
        self.update_preview()

    def rotate_ccw(self):
        self.rotation_angle = (self.rotation_angle - 90) % 360
        self.update_preview()

    def rotate_reset(self):
        self.rotation_angle = 0
        self.update_preview()

    def update_preview(self):
        if self.original_image:
            # Rotate image if needed
            if self.rotation_angle != 0:
                img = self.original_image.rotate(-self.rotation_angle, expand=True)
            else:
                img = self.original_image.copy()

            # Calculate new size
            width, height = img.size
            new_width = int(width * self.preview_scale)
            new_height = int(height * self.preview_scale)

            # Resize image
            img = img.resize((new_width, new_height), Image.LANCZOS)

            # Update the status bar with zoom info
            self.zoom_info.set(f"Zoom: {int(self.preview_scale * 100)}%")

            # Convert to PhotoImage and keep a reference
            self.preview_image = ImageTk.PhotoImage(img)

            # Clear canvas and display image
            self.canvas.delete("all")
            self.canvas.configure(scrollregion=(0, 0, new_width, new_height))
            self.canvas.create_image(0, 0, image=self.preview_image, anchor=tk.NW)

    # File and directory browsing methods
    def browse_input_image(self):
        filetypes = (
            ('Image files', '*.png *.jpg *.jpeg *.bmp *.tiff *.tif'),
            ('All files', '*.*')
        )
        filename = filedialog.askopenfilename(
            title='Open an image file',
            initialdir='/',
            filetypes=filetypes
        )

        if filename:
            self.input_image_path.set(filename)
            self.load_image(filename)

            # Auto-generate output filename
            base_name = os.path.splitext(os.path.basename(filename))[0]
            self.output_doc_path.set(os.path.join(os.path.dirname(filename), f"{base_name}.docx"))

    def load_image(self, image_path):
        try:
            # Open the image and store original
            self.original_image = Image.open(image_path)

            # Reset zoom and rotation
            self.preview_scale = 1.0
            self.rotation_angle = 0

            # Display the image
            self.update_preview()
            self.zoom_fit()  # Auto fit the image

            self.status_message.set(f"Loaded image: {os.path.basename(image_path)}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load image: {str(e)}")
            self.status_message.set("Error loading image")

    def browse_output_file(self):
        filetypes = (
            ('Word documents', '*.docx'),
            ('All files', '*.*')
        )
        filename = filedialog.asksaveasfilename(
            title='Save as Word document',
            initialdir='/',
            defaultextension=".docx",
            filetypes=filetypes
        )

        if filename:
            self.output_doc_path.set(filename)

    def browse_output_dir(self):
        directory = filedialog.askdirectory(
            title='Select output directory for batch processing'
        )

        if directory:
            self.batch_output_dir.set(directory)

    # Batch processing methods
    def add_batch_files(self):
        filetypes = (
            ('Image files', '*.png *.jpg *.jpeg *.bmp *.tiff *.tif'),
            ('All files', '*.*')
        )
        filenames = filedialog.askopenfilenames(
            title='Select image files for batch processing',
            initialdir='/',
            filetypes=filetypes
        )

        if filenames:
            for file in filenames:
                if file not in self.batch_files:
                    self.batch_files.append(file)
                    self.files_listbox.insert(tk.END, os.path.basename(file))

            self.status_message.set(f"Added {len(filenames)} files to batch queue")

    def remove_selected_files(self):
        selected_indices = self.files_listbox.curselection()
        if not selected_indices:
            return

        # Convert to list and sort in reverse order to avoid index issues when deleting
        indices = sorted(list(selected_indices), reverse=True)

        for i in indices:
            del self.batch_files[i]
            self.files_listbox.delete(i)

        self.status_message.set(f"Removed {len(indices)} files from batch queue")

    def clear_batch_files(self):
        self.batch_files.clear()
        self.files_listbox.delete(0, tk.END)
        self.status_message.set("Cleared batch queue")

    def process_batch(self):
        if not self.batch_files:
            messagebox.showerror("Error", "Batch queue is empty. Please add files first.")
            return

        if not self.batch_output_dir.get():
            messagebox.showerror("Error", "Please select an output directory for batch processing.")
            return

        # Start batch processing in a separate thread
        self.progress_bar.start()
        self.status_message.set("Processing batch...")
        threading.Thread(target=self.batch_process_thread, daemon=True).start()

    def batch_process_thread(self):
        success_count = 0
        fail_count = 0
        total_files = len(self.batch_files)

        try:
            for i, image_path in enumerate(self.batch_files):
                try:
                    # Update status on main thread
                    self.root.after(0, lambda: self.status_message.set(
                        f"Processing file {i + 1} of {total_files}: {os.path.basename(image_path)}"))

                    # Load and process the image
                    image = Image.open(image_path)

                    # Apply image processing
                    processed_image = self.process_image_with_settings(image)

                    # Perform OCR
                    text = pytesseract.image_to_string(processed_image, lang=self.language.get())

                    # Generate output filename
                    base_name = os.path.splitext(os.path.basename(image_path))[0]
                    output_file = os.path.join(self.batch_output_dir.get(), f"{base_name}.docx")

                    # Create Word document
                    doc = self.create_word_document(text)

                    # Save document
                    doc.save(output_file)
                    success_count += 1

                except Exception as e:
                    fail_count += 1
                    self.root.after(0, lambda msg=str(e): messagebox.showerror("Error",
                                                                               f"Failed to process {os.path.basename(image_path)}: {msg}"))

            # Complete
            self.root.after(0, self.batch_process_complete, success_count, fail_count)

        except Exception as e:
            self.root.after(0, lambda: messagebox.showerror("Error", f"Batch processing failed: {str(e)}"))
            self.root.after(0, lambda: self.progress_bar.stop())
            self.root.after(0, lambda: self.status_message.set("Batch processing failed"))

    # def batch_process_complete(self, success_count, fail_count):
    #     self.progress_bar.stop()
    #     message = f"Batch processing complete. Success: {success_count}, Failed: {fail_count}"
    #     self.status_message.set(message)
    #     messagebox.showinfo("Batch Complete",

    def batch_process_complete(self, success_count, fail_count):
        self.progress_bar.stop()
        message = f"Batch processing complete. Success: {success_count}, Failed: {fail_count}"
        self.status_message.set(message)
        messagebox.showinfo("Batch Complete", message)

        # Ask if user wants to open the output directory
        if success_count > 0 and messagebox.askyesno("Open Directory",
                                                     "Would you like to open the output directory?"):
            self.open_directory(self.batch_output_dir.get())

    def open_directory(self, dir_path):
        try:
            import platform
            import subprocess

            if platform.system() == 'Windows':
                os.startfile(dir_path)
            elif platform.system() == 'Darwin':  # macOS
                subprocess.call(('open', dir_path))
            else:  # Linux
                subprocess.call(('xdg-open', dir_path))
        except Exception as e:
            messagebox.showerror("Error", f"Could not open directory: {str(e)}")

        # Language handling methods

    def on_language_selected(self, event=None):
        language_name = event.widget.get()
        language_code = self.languages[language_name]
        self.language.set(language_code)

        # Image processing methods

    def apply_image_processing(self):
        if not self.original_image:
            messagebox.showerror("Error", "No image loaded. Please load an image first.")
            return

        try:
            # Process the image
            processed_image = self.process_image_with_settings(self.original_image)

            # Update the display
            self.original_image = processed_image
            self.update_preview()

            self.status_message.set("Image processing applied")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to apply image processing: {str(e)}")

    def process_image_with_settings(self, image):
        img = image.copy()

        # Apply brightness adjustment
        if self.brightness.get() != 1.0:
            enhancer = ImageEnhance.Brightness(img)
            img = enhancer.enhance(self.brightness.get())

        # Apply contrast adjustment
        if self.contrast.get() != 1.0:
            enhancer = ImageEnhance.Contrast(img)
            img = enhancer.enhance(self.contrast.get())

        # Apply sharpness adjustment
        if self.sharpen.get() != 1.0:
            enhancer = ImageEnhance.Sharpness(img)
            img = enhancer.enhance(self.sharpen.get())

        # Apply binarization (convert to black and white)
        if self.binarize.get():
            # Convert to grayscale first
            img = img.convert('L')
            # Apply threshold
            img = img.point(lambda x: 0 if x < self.threshold.get() else 255, '1')

        return img

    def reset_image_processing(self):
        # Reset all image processing values to defaults
        self.brightness.set(1.0)
        self.contrast.set(1.0)
        self.sharpen.set(1.0)
        self.binarize.set(False)
        self.threshold.set(127)

        # If an image is loaded, reset to original
        if hasattr(self, 'input_image_path') and self.input_image_path.get():
            self.load_image(self.input_image_path.get())

        self.status_message.set("Image processing reset to defaults")

        # Text preview methods

    def preview_text(self):
        if not self.original_image:
            messagebox.showerror("Error", "No image loaded. Please load an image first.")
            return

        self.progress_bar.start()
        self.status_message.set("Extracting text...")
        threading.Thread(target=self.extract_text_thread, daemon=True).start()

    def extract_text_thread(self):
        try:
            # Process the image
            processed_image = self.process_image_with_settings(self.original_image)

            # Perform OCR
            text = pytesseract.image_to_string(processed_image, lang=self.language.get())

            # Show text in preview window
            self.root.after(0, lambda: self.show_text_preview(text))

        except Exception as e:
            self.root.after(0, lambda: messagebox.showerror("Error", f"Failed to extract text: {str(e)}"))
            self.root.after(0, lambda: self.progress_bar.stop())
            self.root.after(0, lambda: self.status_message.set("Text extraction failed"))

    def show_text_preview(self, text):
        self.progress_bar.stop()
        self.status_message.set("Text extracted")

        # Create a new window for text preview if it doesn't exist
        if not self.text_preview_window or not tk.Toplevel.winfo_exists(self.text_preview_window):
            self.text_preview_window = tk.Toplevel(self.root)
            self.text_preview_window.title("Extracted Text Preview")
            self.text_preview_window.geometry("600x400")

            # Add controls
            control_frame = ttk.Frame(self.text_preview_window)
            control_frame.pack(fill=tk.X, padx=10, pady=5)

            ttk.Button(control_frame, text="Copy to Clipboard",
                       command=lambda: self.copy_to_clipboard(text_area.get("1.0", tk.END))).pack(side=tk.LEFT,
                                                                                                  padx=5)
            ttk.Button(control_frame, text="Save to Text File",
                       command=lambda: self.save_text_to_file(text_area.get("1.0", tk.END))).pack(side=tk.LEFT,
                                                                                                  padx=5)
            ttk.Button(control_frame, text="Close",
                       command=self.text_preview_window.destroy).pack(side=tk.RIGHT, padx=5)

            # Add text area with scrollbar
            text_frame = ttk.Frame(self.text_preview_window)
            text_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

            text_area = scrolledtext.ScrolledText(text_frame, wrap=tk.WORD, width=80, height=20)
            text_area.pack(fill=tk.BOTH, expand=True)

            # Insert the text
            text_area.delete("1.0", tk.END)
            text_area.insert("1.0", text)
        else:
            # Update existing window
            for widget in self.text_preview_window.winfo_children():
                if isinstance(widget, ttk.Frame):
                    for child in widget.winfo_children():
                        if isinstance(child, scrolledtext.ScrolledText):
                            child.delete("1.0", tk.END)
                            child.insert("1.0", text)
                            break

        # Bring window to front
        self.text_preview_window.lift()
        self.text_preview_window.focus_force()

    def copy_to_clipboard(self, text):
        self.root.clipboard_clear()
        self.root.clipboard_append(text)
        self.status_message.set("Text copied to clipboard")

    def save_text_to_file(self, text):
        filetypes = (
            ('Text files', '*.txt'),
            ('All files', '*.*')
        )
        filename = filedialog.asksaveasfilename(
            title='Save as text file',
            initialdir='/',
            defaultextension=".txt",
            filetypes=filetypes
        )

        if filename:
            try:
                with open(filename, 'w', encoding='utf-8') as file:
                    file.write(text)
                self.status_message.set(f"Text saved to {os.path.basename(filename)}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save text file: {str(e)}")

        # Word document creation and processing methods

    def process_image(self):
        # Check if batch mode is enabled
        if self.batch_mode.get():
            self.process_batch()
            return

        # Validate input
        if not self.input_image_path.get():
            messagebox.showerror("Error", "Please select an input image.")
            return

        if not self.output_doc_path.get():
            messagebox.showerror("Error", "Please specify an output document path.")
            return

        # Start processing in a separate thread
        self.progress_bar.start()
        self.status_message.set("Processing...")
        threading.Thread(target=self.ocr_to_word_thread, daemon=True).start()

    def ocr_to_word_thread(self):
        try:
            image_path = self.input_image_path.get()
            output_file = self.output_doc_path.get()

            # Load the image
            image = Image.open(image_path)

            # Apply image processing
            processed_image = self.process_image_with_settings(image)

            # Extract text using pytesseract OCR
            text = pytesseract.image_to_string(processed_image, lang=self.language.get())

            # Create a Word document
            doc = self.create_word_document(text)

            # Save the document
            doc.save(output_file)

            # Update UI on the main thread
            self.root.after(0, self.process_complete, True,
                            f"Document successfully created: {os.path.basename(output_file)}")

        except Exception as e:
            self.root.after(0, self.process_complete, False, str(e))

    def create_word_document(self, text):
        # Create a new Word document
        doc = Document()

        # Apply document title if enabled
        if self.include_title.get():
            doc.add_heading(self.title_text.get(), 0)

        # Get paragraph alignment
        align_map = {
            "Left": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "Center": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "Right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            "Justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        }
        alignment = align_map.get(self.alignment.get(), WD_PARAGRAPH_ALIGNMENT.LEFT)

        # Process the text content
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph(para.strip())

                # Apply paragraph formatting
                p.alignment = alignment

                # Apply character formatting to runs
                for run in p.runs:
                    run.font.name = self.font_family.get()
                    run.font.size = Pt(self.font_size.get())

        return doc

    def process_complete(self, success, message):
        self.progress_bar.stop()

        if success:
            self.status_message.set(message)
            messagebox.showinfo("Success", message)

            # Ask if user wants to open the document
            if messagebox.askyesno("Open Document", "Would you like to open the created document?"):
                self.open_document(self.output_doc_path.get())
        else:
            self.status_message.set(f"Error: {message}")
            messagebox.showerror("Error", f"Failed to convert image: {message}")

    def open_document(self, doc_path):
        try:
            import platform
            import subprocess

            if platform.system() == 'Windows':
                os.startfile(doc_path)
            elif platform.system() == 'Darwin':  # macOS
                subprocess.call(('open', doc_path))
            else:  # Linux
                subprocess.call(('xdg-open', doc_path))
        except Exception as e:
            messagebox.showerror("Error", f"Could not open document: {str(e)}")

def main():
    root = tk.Tk()
    app = OCRtoWordGUI(root)
    root.mainloop()

if __name__ == "__main__":
    main()